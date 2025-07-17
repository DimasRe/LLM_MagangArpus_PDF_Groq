# Migrated FastAPI backend using SQLAlchemy + MySQL (XAMPP)

from fastapi import FastAPI, HTTPException, Body, Query, UploadFile, File, Form, Depends
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
from dotenv import load_dotenv
from sqlalchemy import create_engine, Column, String, Integer, Text, Boolean, DateTime
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
from datetime import datetime
import os
import uuid
import json
import requests
import shutil
from pathlib import Path
import docx
from PyPDF2 import PdfReader

# Load environment variables
load_dotenv()

# --- Configuration Constants ---
UPLOAD_DIR = os.getenv("UPLOAD_DIR", "uploads")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
DATABASE_URL = os.getenv("DATABASE_URL") # Pastikan ini diatur di .env Anda!

# Ensure uploads directory exists
Path(UPLOAD_DIR).mkdir(exist_ok=True)

# Check if GROQ API key is provided and warn if not
if not GROQ_API_KEY:
    print("⚠️  PERINGATAN: GROQ_API_KEY tidak ditemukan di variabel lingkungan!")
    print("   Harap buat file .env dengan kunci API GROQ Anda")
    print("   Dapatkan kunci API gratis Anda di: https://console.groq.com/")

# Check if DATABASE_URL is provided and warn if not
if not DATABASE_URL:
    print("⚠️  PERINGATAN: DATABASE_URL tidak ditemukan di variabel lingkungan!")
    print("   Harap buat file .env dengan URL database MySQL Anda (contoh: mysql+pymysql://user:password@host:3306/database_name)")

# --- Setup DB engine (SQLAlchemy) ---
# Menggunakan pool_recycle untuk mengatasi masalah "MySQL has gone away"
engine = create_engine(DATABASE_URL, pool_recycle=3600)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

# --- Database Models (SQLAlchemy ORM) ---
class Document(Base):
    __tablename__ = "documents"
    id = Column(String(36), primary_key=True, index=True) # UUIDs are 36 chars
    filename = Column(String(255), nullable=False)
    file_path = Column(String(512), nullable=False)
    upload_date = Column(DateTime, nullable=False)
    text_content = Column(Text) # Text for potentially large content
    file_size = Column(Integer)

class ChatHistory(Base):
    __tablename__ = "chat_history"
    id = Column(Integer, primary_key=True, index=True, autoincrement=True)
    session_id = Column(String(255), nullable=False)
    message = Column(Text, nullable=False)
    response = Column(Text, nullable=False)
    timestamp = Column(DateTime, nullable=False)
    is_predefined = Column(Boolean, default=False)
    document_ids = Column(Text) # Storing JSON string of list of IDs

# Create tables in the database (if they don't exist)
Base.metadata.create_all(bind=engine)

# --- FastAPI App Initialization ---
app = FastAPI(
    title="Chatbot Dinas Arpus Jateng",
    description="Sistem Analisis Dokumen dan Chat untuk Dinas Kearsipan dan Perpustakaan Provinsi Jawa Tengah (Akses Publik)",
    version="2.0.0"
)

# Add CORS middleware for development
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # For development; restrict in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Pydantic Models for API Data Validation ---
class ChatMessage(BaseModel):
    """Represents an incoming chat message."""
    message: str
    document_ids: List[str] = []
    is_predefined: bool = False

class ChatResponse(BaseModel):
    """Represents the response from the chatbot."""
    response: str
    source_documents: List[str] = []
    predefined_questions: List[str] = []

class SystemHealth(BaseModel):
    """Represents the health status of the system."""
    status: str
    groq_api: str
    database: str
    ai_info: Optional[Dict[str, Any]] = None

class AdminStats(BaseModel):
    """Represents statistics for the admin dashboard."""
    total_documents: int
    total_chats: int
    recent_activity: List[Dict[str, Any]]

# --- Predefined Questions ---
PREDEFINED_QUESTIONS = [
    "Apa ringkasan dari dokumen ini?",
    "Kapan dokumen ini dibuat dan oleh siapa?",
    "Apa tujuan utama dokumen ini?",
    "Bagaimana relevansi dokumen ini dengan Dinas Arpus Jateng?",
    "Informasi penting apa yang ada dalam dokumen ini?",
    "Bagian mana dari dokumen ini yang membahas tentang [topik spesifik]?",
    "Bisakah Anda menemukan data atau angka penting di dokumen ini?",
    "Apakah ada rekomendasi atau kebijakan yang disebutkan dalam dokumen ini?",
    "Apa kesimpulan dari dokumen ini?"
]

# --- Database Dependency (for FastAPI) ---
def get_db():
    """Dependency to get a SQLAlchemy session for a request."""
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# --- Document Text Extraction Functions ---
def extract_text_from_pdf(file_path: str) -> str:
    """Extracts text content from a PDF file."""
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        print(f"Error extracting text from PDF '{file_path}': {e}")
        return ""

def extract_text_from_docx(file_path: str) -> str:
    """Extracts text content from a DOCX file."""
    try:
        doc = docx.Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        print(f"Error extracting text from DOCX '{file_path}': {e}")
        return ""

def extract_text_from_file(file_path: str) -> str:
    """
    Extracts text content based on file extension.
    Supports PDF, DOCX, DOC, and TXT files.
    """
    file_extension = Path(file_path).suffix.lower().lstrip('.')

    if file_extension == "pdf":
        return extract_text_from_pdf(file_path)
    elif file_extension in ["docx", "doc"]:
        return extract_text_from_docx(file_path)
    elif file_extension == "txt":
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            print(f"Error reading TXT file '{file_path}': {e}")
            return ""
    else:
        print(f"Unsupported file type for extraction: {file_extension}")
        return ""

# --- GROQ AI Interaction Functions ---
def query_groq(prompt: str, max_tokens: int = 2000, model: str = "llama3-8b-8192") -> str:
    """
    Queries the GROQ API for AI responses.

    Args:
        prompt (str): The text prompt to send to the AI.
        max_tokens (int): The maximum number of tokens to generate in the response.
        model (str): The AI model to use.

    Returns:
        str: The AI's response, or an error message if the query fails.
    """
    if not GROQ_API_KEY:
        return "Error: GROQ API key not configured. Please check your .env file."

    try:
        headers = {
            "Authorization": f"Bearer {GROQ_API_KEY}",
            "Content-Type": "application/json"
        }

        payload = {
            "model": model,
            "messages": [
                {
                    "role": "system",
                    "content": "Anda adalah asisten analisis dokumen yang membantu staf dan publik di Dinas Kearsipan dan Perpustakaan Provinsi Jawa Tengah (Dinas Arpus Jateng). Berikan jawaban yang akurat, informatif, dan relevan dalam bahasa Indonesia."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "max_tokens": max_tokens,
            "temperature": 0.7,
            "top_p": 0.9,
            "stream": False
        }

        response = requests.post(GROQ_API_URL, json=payload, headers=headers, timeout=30)
        response.raise_for_status() # Raise HTTPError for bad responses (4xx or 5xx)

        result = response.json()
        if "choices" in result and len(result["choices"]) > 0:
            return result["choices"][0]["message"]["content"]
        else:
            return "Error: Invalid response format from GROQ API."

    except requests.exceptions.ConnectionError:
        return "Error: Tidak dapat terhubung ke GROQ API. Periksa koneksi internet Anda."
    except requests.exceptions.Timeout:
        return "Error: Permintaan ke GROQ API habis waktu (timeout). Silakan coba lagi."
    except requests.exceptions.HTTPError as e:
        status_code = e.response.status_code
        if status_code == 401:
            return "Error: Kunci API GROQ tidak valid. Periksa kredensial Anda."
        elif status_code == 429:
            return "Error: Batas permintaan (rate limit) terlampaui. Silakan coba lagi nanti."
        else:
            print(f"GROQ API HTTP error: {status_code} - {e.response.text}")
            return f"Error: GROQ API mengembalikan status {status_code}."
    except Exception as e:
        print(f"Error querying GROQ: {e}")
        return f"Error internal saat berinteraksi dengan AI: {str(e)}"

def test_groq_connection() -> Dict[str, str]:
    """Tests the GROQ API connection and returns its status."""
    if not GROQ_API_KEY:
        return {"status": "error", "message": "API key not configured."}

    try:
        # Use a short, specific prompt for connection test
        response = query_groq("Test koneksi, balas 'Koneksi berhasil.'", max_tokens=20, model="llama3-8b-8192")
        if "koneksi berhasil" in response.lower():
            return {
                "status": "connected",
                "message": "GROQ API berfungsi dengan baik.",
                "model": "llama3-8b-8192"
            }
        else:
            return {"status": "error", "message": f"Respon tidak terduga dari GROQ: {response[:100]}..."}
    except Exception as e:
        return {"status": "error", "message": str(e)}

# --- Dummy Admin Authentication (FOR DEMONSTRATION ONLY!) ---
def get_admin_status_dummy(is_admin_query: bool = Query(False, description="Set to true to access admin features. FOR DEMO ONLY!")):
    """
    Dummy function to simulate admin access via query parameter.
    **DO NOT USE IN PRODUCTION WITHOUT PROPER AUTHENTICATION!**
    """
    if is_admin_query:
        return {"username": "admin_demo", "role": "admin"}
    raise HTTPException(status_code=403, detail="Akses admin diperlukan (gunakan ?is_admin_query=true untuk demo)")


# --- API Endpoints ---

@app.get("/health", response_model=SystemHealth, tags=["System"])
def health_check(db: SessionLocal = Depends(get_db)):
    """
    Checks the health of the API and its dependencies (GROQ API, Database).
    """
    health_status = {
        "status": "healthy",
        "groq_api": "disconnected",
        "database": "disconnected",
        "ai_info": None
    }

    # Check GROQ API connection
    groq_test = test_groq_connection()
    health_status["groq_api"] = groq_test["status"]
    if groq_test["status"] == "connected":
        health_status["ai_info"] = {
            "provider": "GROQ",
            "model": groq_test.get("model", "N/A"),
            "status": "operasional"
        }
    else:
        # If GROQ is not connected, provide the specific error message
        health_status["ai_info"] = {
            "provider": "GROQ",
            "model": groq_test.get("model", "N/A"),
            "status": "non-operasional",
            "error": groq_test.get("message", "Unknown error")
        }


    # Check database connection
    try:
        db.execute("SELECT 1") # Simple query to test connection
        health_status["database"] = "connected"
    except Exception as e:
        health_status["database"] = "disconnected"
        print(f"Database connection error: {e}")

    # Overall status
    if health_status["groq_api"] == "connected" and health_status["database"] == "connected":
        health_status["status"] = "healthy"
    else:
        health_status["status"] = "degraded" # At least one dependency is not healthy

    return health_status

@app.post("/upload", tags=["Documents"])
async def upload_documents(
    files: List[UploadFile] = File(...),
    db: SessionLocal = Depends(get_db) # Inject DB session
):
    """
    Uploads documents for processing.

    Args:
        files (List[UploadFile]): A list of files to upload.

    Returns:
        Dict: A dictionary containing information about uploaded documents.

    Raises:
        HTTPException: If more than 5 files are uploaded or file type is not supported.
    """
    if len(files) > 5:
        raise HTTPException(status_code=400, detail="Maksimal 5 file diizinkan.")

    uploaded_docs = []
    for file in files:
        # Generate a unique ID for the document
        doc_id = str(uuid.uuid4())
        file_extension = Path(file.filename).suffix.lower().lstrip('.')

        # Validate file extension
        if file_extension not in ["pdf", "docx", "doc", "txt"]:
            print(f"Skipping unsupported file type: {file.filename}")
            continue # Skip to next file if unsupported

        file_path = os.path.join(UPLOAD_DIR, f"{doc_id}.{file_extension}")

        # Save the file to disk
        try:
            with open(file_path, "wb") as f:
                shutil.copyfileobj(file.file, f)
        except Exception as e:
            print(f"Error saving file '{file.filename}': {e}")
            raise HTTPException(status_code=500, detail=f"Gagal menyimpan file '{file.filename}'.")

        # Extract text content for AI processing and save to DB
        text = extract_text_from_file(file_path)
        file_size = os.path.getsize(file_path) # Get actual file size

        try:
            doc = Document(
                id=doc_id,
                filename=file.filename,
                file_path=file_path,
                upload_date=datetime.now(),
                text_content=text,
                file_size=file_size
            )
            db.add(doc)
            db.commit() # Commit each document individually or batch later
            db.refresh(doc) # Refresh to get any DB-generated values if needed

            uploaded_docs.append({
                "document_id": doc_id,
                "filename": file.filename,
                "size": file_size, # Use actual file size
                "predefined_questions": PREDEFINED_QUESTIONS # Suggest predefined questions
            })
        except Exception as e:
            db.rollback()
            print(f"Database error saving document '{file.filename}': {e}")
            raise HTTPException(status_code=500, detail=f"Gagal menyimpan metadata dokumen '{file.filename}'.")

    if not uploaded_docs and files:
        raise HTTPException(status_code=400, detail="Tidak ada file yang diunggah karena format tidak didukung.")
    elif not uploaded_docs and not files:
        raise HTTPException(status_code=400, detail="Tidak ada file yang dipilih untuk diunggah.")

    return {"uploaded_documents": uploaded_docs, "message": f"{len(uploaded_docs)} dokumen berhasil diunggah."}


@app.get("/predefined-questions/{document_id}", tags=["Chat"])
def get_predefined_questions(
    document_id: str,
    db: SessionLocal = Depends(get_db)
):
    """
    Gets predefined questions relevant to a specific document.

    Args:
        document_id (str): The ID of the document to get questions for.

    Returns:
        Dict: A dictionary containing a list of predefined questions.

    Raises:
        HTTPException: If the document is not found.
    """
    doc = db.query(Document).filter(Document.id == document_id).first()

    if not doc:
        raise HTTPException(status_code=404, detail="Dokumen tidak ditemukan.")

    return {"questions": PREDEFINED_QUESTIONS, "document_id": document_id}

@app.post("/chat", response_model=ChatResponse, tags=["Chat"])
def chat(
    message: ChatMessage,
    db: SessionLocal = Depends(get_db)
):
    """
    Handles chat interactions, responding based on provided documents or general knowledge.

    Args:
        message (ChatMessage): The incoming chat message including the query and document IDs.

    Returns:
        ChatResponse: The AI's response, source documents, and potentially predefined questions.
    """
    session_id = "guest_session" # Using a fixed session ID as no user login is implemented

    document_texts = []
    document_names = []
    if message.document_ids:
        # Fetch documents using SQLAlchemy
        docs = db.query(Document).filter(Document.id.in_(message.document_ids)).all()

        for doc in docs:
            if doc.text_content:
                # IMPORTANT: Consider the token limit of the AI model.
                # Llama 3 8B has an 8192 token context window.
                # 4000 characters is roughly ~1000 tokens for English.
                # If documents are large, consider more sophisticated chunking/retrieval methods
                # or increasing this limit if the model supports it.
                document_texts.append(doc.text_content[:4000]) # Keep existing truncation
                document_names.append(doc.filename)
            else:
                print(f"Warning: Document ID {doc.id} has no text content for chat context.")

        # Check if any requested documents were not found
        found_doc_ids = {doc.id for doc in docs}
        not_found_ids = [doc_id for doc_id in message.document_ids if doc_id not in found_doc_ids]
        if not_found_ids:
            print(f"Warning: Document IDs {not_found_ids} not found for chat context.")


    # System prompt for the AI to define its persona and rules
    strict_rules = """
Anda adalah asisten analisis dokumen yang membantu staf dan publik di Dinas Kearsipan dan Perpustakaan Provinsi Jawa Tengah (Dinas Arpus Jateng).
Tugas utama Anda adalah menganalisis dokumen arsip, perpustakaan, dan dokumen resmi lainnya serta menjawab pertanyaan yang relevan.
**ATURAN UTAMA:**
1.  **FOKUS TOPIK:** Anda **HANYA** boleh menjawab pertanyaan yang berkaitan dengan:
    a. Isi dari dokumen yang diberikan.
    b. Informasi umum yang relevan dengan konteks Dinas Kearsipan dan Perpustakaan Provinsi Jawa Tengah (Dinas Arpus Jateng).
2.  **TOLAK PERTANYAAN TIDAK RELEVAN:** Jika pertanyaan pengguna berada di luar dua topik di atas (contoh: pertanyaan tentang cuaca, resep, pemrograman, atau topik umum lainnya), Anda **HARUS** menolak dengan sopan. Gunakan jawaban ini: "Maaf, saya hanya dapat merespons pertanyaan yang berkaitan dengan analisis dokumen atau informasi seputar Dinas Kearsipan dan Perpustakaan Provinsi Jawa Tengah."
3.  **BAHASA:** Selalu gunakan Bahasa Indonesia yang baik dan formal.
---
"""

    prompt = strict_rules

    if document_texts:
        prompt += "Konteks dari dokumen yang disediakan:\n"
        for i, (text_chunk, name) in enumerate(zip(document_texts, document_names)):
            prompt += f"\n--- DOKUMEN {i+1}: {name} ---\n{text_chunk}\n" # Use text_chunk (already truncated)

        prompt += "\nBerdasarkan aturan di atas dan konteks dari dokumen yang disediakan, jawablah pertanyaan berikut.\n"
    else:
        prompt += "\nBerdasarkan aturan di atas dan pengetahuan umum Anda tentang Dinas Kearsipan dan Perpustakaan Provinsi Jawa Tengah, jawablah pertanyaan berikut. Tidak ada dokumen yang disediakan.\n"

    prompt += f"\nPertanyaan Pengguna: \"{message.message}\""

    # Query the GROQ AI
    response_content = query_groq(prompt, max_tokens=1500)

    # Save chat history to database
    try:
        chat_entry = ChatHistory(
            session_id=session_id, # Or generate new UUID if each chat interaction is a new session
            message=message.message,
            response=response_content,
            timestamp=datetime.now(),
            is_predefined=message.is_predefined,
            document_ids=json.dumps(message.document_ids) # Store as JSON string
        )
        db.add(chat_entry)
        db.commit()
        db.refresh(chat_entry)
    except Exception as e:
        db.rollback()
        print(f"Database error saving chat history: {e}")
        # This error is logged but not raised to the user, as the chat response itself is more critical.

    # Determine if predefined questions should be suggested
    predefined_questions_to_suggest = PREDEFINED_QUESTIONS if message.document_ids and not message.is_predefined else []

    return {
        "response": response_content,
        "source_documents": document_names,
        "predefined_questions": predefined_questions_to_suggest
    }

@app.get("/documents", tags=["Documents"])
def get_documents(db: SessionLocal = Depends(get_db)):
    """
    Retrieves a list of all available documents. (Public access)

    Returns:
        Dict: A dictionary containing a list of document metadata.
    """
    docs = db.query(Document).order_by(Document.upload_date.desc()).all()

    # Convert Document objects to dictionaries for JSON serialization
    return {"documents": [
        {
            "id": d.id,
            "filename": d.filename,
            "upload_date": d.upload_date.isoformat(), # Convert datetime to ISO string
            "file_size": d.file_size
        } for d in docs
    ]}

@app.get("/history", tags=["Chat"])
def get_chat_history(db: SessionLocal = Depends(get_db)):
    """
    Retrieves recent chat history. (Public access, simplified)

    Returns:
        Dict: A dictionary containing a list of chat history items.
    """
    history_records = db.query(ChatHistory).order_by(ChatHistory.timestamp.desc()).limit(100).all()

    # Convert ChatHistory objects to dictionaries for JSON serialization
    # And parse document_ids back from JSON string if needed, or keep as string for now
    return {"history": [
        dict(
            id=item.id,
            session_id=item.session_id,
            message=item.message,
            response=item.response,
            timestamp=item.timestamp.isoformat(), # Convert datetime to ISO string
            is_predefined=item.is_predefined,
            document_ids=json.loads(item.document_ids) if item.document_ids else [], # Parse JSON
            username=item.session_id # For frontend display convenience
        ) for item in history_records
    ]}

@app.delete("/history/{history_id}", tags=["Chat"])
def delete_chat_history(history_id: int, db: SessionLocal = Depends(get_db)):
    """
    Deletes a specific chat history item by ID.

    Args:
        history_id (int): The ID of the chat history item to delete.

    Returns:
        Dict: A success message.

    Raises:
        HTTPException: If the history item is not found or deletion fails.
    """
    try:
        chat_item = db.query(ChatHistory).filter(ChatHistory.id == history_id).first()
        if not chat_item:
            raise HTTPException(status_code=404, detail="Riwayat chat tidak ditemukan.")

        db.delete(chat_item)
        db.commit()

        return {"message": "Riwayat chat berhasil dihapus.", "success": True}

    except HTTPException:
        db.rollback()
        raise
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Gagal menghapus riwayat: {str(e)}")


@app.delete("/history", tags=["Chat"])
def delete_all_chat_history(db: SessionLocal = Depends(get_db)):
    """
    Deletes all chat history entries.

    Returns:
        Dict: A success message with the count of deleted items.

    Raises:
        HTTPException: If deletion fails.
    """
    try:
        deleted_count = db.query(ChatHistory).delete()
        db.commit()

        return {"message": f"Semua riwayat chat berhasil dihapus ({deleted_count} item).", "success": True, "deleted_count": deleted_count}

    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Gagal menghapus semua riwayat: {str(e)}")

# --- Admin-only Endpoints (FOR DEMONSTRATION ONLY!) ---
@app.get("/admin/users", tags=["Admin"])
def get_all_users_admin(admin_status: dict = Depends(get_admin_status_dummy)):
    """
    Retrieves a list of dummy users. (Admin only - FOR DEMO)
    """
    return [] # No actual user management in this public access version

@app.delete("/admin/users/{username}", tags=["Admin"])
def delete_user_admin(username: str, admin_status: dict = Depends(get_admin_status_dummy)):
    """
    Dummy endpoint for deleting a user. (Admin only - FOR DEMO)
    """
    raise HTTPException(status_code=400, detail="Manajemen pengguna dinonaktifkan dalam mode akses publik.")

@app.get("/admin/stats", response_model=AdminStats, tags=["Admin"])
def get_admin_stats(admin_status: dict = Depends(get_admin_status_dummy), db: SessionLocal = Depends(get_db)):
    """
    Retrieves system statistics for the admin dashboard. (Admin only)
    """
    try:
        total_documents = db.query(Document).count()
        total_chats = db.query(ChatHistory).count()

        recent_activity = []
        # Fetch recent chats
        recent_chats = db.query(ChatHistory).order_by(ChatHistory.timestamp.desc()).limit(5).all()

        for chat in recent_chats:
            document_info = "Konteks Umum"
            try:
                doc_ids = json.loads(chat.document_ids) if chat.document_ids else []
                if doc_ids:
                    # Fetch filenames for context, if needed for richer display
                    # For simplicity, just list IDs in demo
                    document_info = f"Dok. ID: {', '.join(doc_ids)}"
            except json.JSONDecodeError:
                pass # Fallback to default if document_ids is malformed

            recent_activity.append({
                "type": "chat",
                "username": chat.session_id,
                "description": f"Bertanya ({document_info}): {chat.message[:70]}{'...' if len(chat.message) > 70 else ''}",
                "timestamp": chat.timestamp.isoformat()
            })

        # Fetch recent uploads
        recent_uploads = db.query(Document).order_by(Document.upload_date.desc()).limit(5).all()

        for upload in recent_uploads:
            recent_activity.append({
                "type": "upload",
                "username": "Pengunggah", # Assuming uploads are by the "public user" in this setup
                "description": f"Mengunggah: {upload.filename}",
                "timestamp": upload.upload_date.isoformat()
            })

        # Sort combined activities by timestamp
        recent_activity.sort(key=lambda x: x["timestamp"], reverse=True)
        recent_activity = recent_activity[:10] # Limit to top 10 recent activities

        return AdminStats(
            total_documents=total_documents,
            total_chats=total_chats,
            recent_activity=recent_activity
        )
    except Exception as e:
        print(f"Error getting admin stats: {e}")
        raise HTTPException(status_code=500, detail=f"Gagal memuat statistik admin: {str(e)}")


@app.get("/admin/documents", tags=["Admin"])
def get_all_documents_admin(admin_status: dict = Depends(get_admin_status_dummy), db: SessionLocal = Depends(get_db)):
    """
    Retrieves all documents in the system for administrative view. (Admin only)
    """
    documents = db.query(Document).order_by(Document.upload_date.desc()).all()

    # For admin view, you might want to show who uploaded it, but current schema
    # doesn't link documents to users. Adding dummy username/email for UI compatibility.
    return {"documents": [
        dict(
            id=doc.id,
            filename=doc.filename,
            upload_date=doc.upload_date.isoformat(),
            file_size=doc.file_size,
            username="N/A",
            email="N/A"
        ) for doc in documents
    ]}

@app.delete("/admin/documents/{document_id}", tags=["Admin"])
def delete_document_admin(document_id: str, admin_status: dict = Depends(get_admin_status_dummy), db: SessionLocal = Depends(get_db)):
    """
    Deletes a document by its ID, including its associated file and chat history. (Admin only)

    Args:
        document_id (str): The ID of the document to delete.

    Returns:
        Dict: A success message.

    Raises:
        HTTPException: If the document is not found or deletion fails.
    """
    try:
        # Get file path before deleting from DB
        doc_to_delete = db.query(Document).filter(Document.id == document_id).first()

        if not doc_to_delete:
            raise HTTPException(status_code=404, detail="Dokumen tidak ditemukan.")

        file_path = doc_to_delete.file_path

        # Delete from documents table
        db.delete(doc_to_delete)
        # Delete related chat history entries
        # Note: Filtering JSON text field might be slow on large tables
        db.query(ChatHistory).filter(ChatHistory.document_ids.like(f'%"{document_id}"%')).delete(synchronize_session=False)

        db.commit()

        # Delete the physical file after successful DB operations
        if os.path.exists(file_path):
            os.remove(file_path)
            print(f"Successfully deleted file: {file_path}")
        else:
            print(f"Warning: File not found on disk, but removed from DB: {file_path}")

        return {"message": "Dokumen berhasil dihapus."}
    except HTTPException:
        db.rollback()
        raise
    except Exception as e:
        db.rollback()
        print(f"Error deleting document {document_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Gagal menghapus dokumen: {str(e)}")


# --- Frontend Serving ---
# Serve the main index.html file
@app.get("/", response_class=FileResponse, include_in_schema=False)
async def read_index():
    """Serves the main index.html file."""
    # This assumes index.html is in the same directory as this script.
    # If your index.html is in a subdirectory (e.g., "static"), adjust the path.
    return "index.html"

# Mount static files (CSS, JS, images etc.)
# This will serve files from the current directory.
# If your static files are in a specific folder (e.g., 'static'), change directory="." to directory="static"
app.mount("/", StaticFiles(directory="."), name="static_root")


# --- Main Application Runner ---
if __name__ == "__main__":
    import uvicorn
    print("Memulai Chatbot Dinas Arpus Jateng dengan GROQ AI (Akses Publik)")
    print("Dokumentasi API: http://localhost:8000/docs")
    print("Aplikasi Frontend: http://localhost:8000")
    print("Demo akses admin: http://localhost:8000/?is_admin_query=true (hanya untuk pengujian)")
    uvicorn.run(app, host="0.0.0.0", port=8000)